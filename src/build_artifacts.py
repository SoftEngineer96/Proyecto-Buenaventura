from __future__ import annotations

import json
from pathlib import Path

import nbformat as nbf
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
LANDING = ROOT / "data" / "landing"
PROCESSED = ROOT / "data" / "trusted" / "eda_historico_2012_2024"
FIGURES = ROOT / "reportes" / "figuras"
OUT = ROOT / "reportes"
OUT.mkdir(parents=True, exist_ok=True)
r = json.loads((PROCESSED / "eda_results.json").read_text(encoding="utf-8"))


def usd(value):
    return f"US$ {value:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def num(value, decimals=0):
    return f"{value:,.{decimals}f}".replace(",", "X").replace(".", ",").replace("X", ".")


def q(title, code, finding, implication):
    return {"title": title, "code": code.strip(), "finding": finding, "implication": implication}


questions = [
q("¿Cuántas filas y columnas tiene el conjunto filtrado?", "df.shape",
  f"El conjunto contiene {num(r['rows'])} registros y {r['columns']} columnas después de filtrar la aduana 35.",
  "El volumen exige lectura por bloques en la extracción y un archivo analítico agregado por mes para el entrenamiento."),
q("¿Qué tipos de datos quedaron después de la normalización?", "df.dtypes.astype(str).value_counts()",
  "Las variables cuantitativas quedaron normalizadas como numéricas; fecha se convirtió a datetime y los identificadores se conservaron como códigos.",
  "Los códigos de país, capítulo y transporte deben tratarse como categorías, aunque estén representados con números."),
q("¿Cuánta memoria ocupa el conjunto en memoria?", "df.memory_usage(deep=True).sum() / 1024**2",
  f"La tabla ocupa aproximadamente {num(r['memory_mb'], 2)} MB en pandas.",
  "En Colab conviene cargar el archivo filtrado, seleccionar columnas y construir agregados antes de modelar."),
q("¿Cuál es la granularidad de una fila?", "df[['FECH','ADUA','PAISGEN','NABAN','VACID','PNK']].head()",
  "Cada fila representa un registro o línea de importación asociado a un mes, una aduana, un país y una subpartida arancelaria; no representa un mes completo.",
  "El modelo de pronóstico debe trabajar con una tabla mensual agregada, no con líneas individuales mezcladas aleatoriamente."),
q("¿Cuál es la cobertura temporal y existen meses faltantes?", "monthly.index.min(), monthly.index.max(), len(monthly), monthly.isna().all(axis=1).sum()",
  f"La serie cubre de {r['date_min']} a {r['date_max']}, tiene {r['months']} meses y no presenta meses ausentes.",
  "La continuidad permite usar rezagos de 1 a 12 meses sin reconstruir periodos faltantes."),
q("¿Existe una clave primaria y cuántos duplicados exactos hay?", "business_cols=[c for c in df if not c.startswith('SOURCE_') and c!='fecha']; df.duplicated(business_cols).sum()",
  f"No existe un identificador único de línea. Se encontraron {num(r['duplicates'])} coincidencias exactas ({num(r['duplicates_pct'],3)} %).",
  "No deben eliminarse automáticamente: líneas comerciales distintas pueden compartir todos los campos publicados. Se auditarán por archivo y mes."),

q("¿Qué columnas tienen valores nulos?", "(df.isna().mean().mul(100).sort_values(ascending=False).loc[lambda s:s>0])",
  "Los nulos son bajos: SEGUROS 0,0040 %, FLETE 0,0018 %, VAFODO 0,0001 % y DEREL 0,0001 %.",
  "Se puede imputar cero solo cuando el diccionario confirme ausencia de cargo; para FOB se debe excluir o imputar con relación CIF-FOB."),
q("¿Qué proporción de filas contiene al menos un nulo?", "df.isna().any(axis=1).mean()*100",
  f"Solo {num(r['rows_with_null_pct'],3)} % de las filas contiene al menos un valor nulo.",
  "La eliminación puntual tendría poco impacto, pero la decisión debe hacerse por variable y significado, no de forma global."),
q("¿El patrón de nulos puede clasificarse como MCAR, MAR o MNAR?", "df.isna().groupby(df['fecha'].dt.year).mean().max().sub(df.isna().groupby(df['fecha'].dt.year).mean().min()).sort_values(ascending=False).head()",
  "No es posible probar MNAR con estos datos. La mayor variación anual de ausencia fue 0,031 puntos porcentuales en SEGUROS, por lo que el patrón es pequeño pero puede depender del periodo de reporte.",
  "Se tratará como potencialmente MAR por año y se incluirán indicadores de ausencia si la imputación mejora la validación temporal."),
q("¿Existen columnas constantes?", "[c for c in df if df[c].nunique(dropna=False)<=1]",
  "ADUA es constante porque el conjunto ya fue filtrado al código 35, Buenaventura.",
  "ADUA debe retirarse del conjunto de features: no aporta variación predictiva."),
q("¿Cuál es la cardinalidad de las variables principales?", "df[['PAISGEN','capitulo','VIATRANS','NABAN']].nunique()",
  f"Hay {r['countries']} países de origen, {r['chapters']} capítulos, {r['transport_modes']} modos de transporte y una cardinalidad mucho mayor en subpartida.",
  "Se debe controlar la dimensionalidad con capítulos, top-N y agrupación 'Otros', en vez de one-hot de todas las subpartidas."),
q("¿Hubo errores de tipado al consolidar los años?", "df[['FECH','VACID','PNK','VAFODO']].dtypes",
  "Los archivos alternaban coma, punto y punto y coma, y algunos meses usaban punto de miles con coma decimal. Después de la normalización no quedaron conversiones numéricas inválidas.",
  "La función de ingestión debe conservar la detección de separador y la conversión localizada para futuras actualizaciones."),
q("¿Hay categorías mal escritas o inconsistentes?", "df[['PAISGEN','capitulo','VIATRANS']].astype('Int64').nunique()",
  "Las categorías analíticas son códigos, por lo que no se detectan faltas ortográficas en ellas. Sí hubo nombres de archivos y rutas mensuales inconsistentes.",
  "Las etiquetas legibles deben incorporarse desde catálogos oficiales mediante joins; no deben inferirse desde nombres de archivos."),
q("¿Existen valores imposibles?", "{c:int((df[c]<0).sum()) for c in ['PNK','VACID','VADUA','BASEIVA','DEREL']}",
  "No hay pesos ni valores CIF negativos. Se encontró un VADUA negativo y una BASEIVA negativa, casos aislados que requieren revisión documental.",
  "Esos dos registros deben marcarse y excluirse de features monetarias derivadas si no corresponden a ajustes válidos."),

q("¿Cuál será la variable objetivo y el horizonte de predicción?", "monthly['cif_usd']=df.groupby('fecha')['VACID'].sum(); target=monthly['cif_usd']",
  "El objetivo principal es el valor CIF mensual en dólares, calculado como suma de VACID. El horizonte propuesto es un mes adelante.",
  "Cada fila de entrenamiento representará un mes y solo usará información disponible hasta el cierre del mes anterior."),
q("¿Cómo se distribuye el target mensual?", "monthly['cif_usd'].describe()",
  f"La media es {usd(r['target_mean'])}, la mediana {usd(r['target_median'])}, el mínimo {usd(r['target_min'])} y el máximo {usd(r['target_max'])}.",
  "Se debe evaluar con métricas absolutas y porcentuales para que los meses de mayor valor no dominen toda la evaluación."),
q("¿El target presenta asimetría o curtosis?", "monthly['cif_usd'].agg(['skew','kurt'])",
  f"La asimetría es {num(r['target_skew'],4)} y la curtosis excedente {num(r['target_kurtosis'],4)}.",
  "La distribución es sesgada a la derecha; conviene comparar entrenamiento en escala original y logarítmica."),
q("¿Qué meses son outliers del target?", "s=monthly['cif_usd']; q1,q3=s.quantile([.25,.75]); iqr=q3-q1; s[(s<q1-1.5*iqr)|(s>q3+1.5*iqr)]",
  "El criterio IQR identifica ocho meses, concentrados entre diciembre de 2021 y septiembre de 2022; z-score mayor que 3 identifica marzo y agosto de 2022.",
  "No se eliminarán: son periodos reales de expansión. Se usarán modelos robustos, indicadores de régimen y análisis de error por periodo."),
q("¿Hay meses con target nulo?", "monthly['cif_usd'].isna().sum()",
  f"El target tiene {r['target_nulls']} meses nulos.",
  "No se requiere imputación temporal del objetivo; cualquier mes nuevo incompleto deberá bloquearse antes de inferencia."),
q("¿Qué variables provocarían fuga de información?", "monthly.corr(method='spearman')['cif_usd'].sort_values(ascending=False)",
  "FOB contemporáneo correlaciona 0,9971 con CIF y flete contemporáneo 0,9066. Ambos forman parte del mismo mes que se intenta predecir.",
  "Solo se usarán sus rezagos. Utilizar valores contemporáneos produciría data leakage y una precisión irreal."),
q("¿Una transformación logarítmica mejora la forma del target?", "monthly['cif_usd'].skew(), np.log1p(monthly['cif_usd']).skew()",
  f"La asimetría baja de {num(r['target_skew'],4)} a {num(r['log_skew'],4)} con log1p.",
  "Ridge y modelos de boosting deben probarse sobre log1p(CIF), revirtiendo con expm1 y recortando predicciones negativas."),
q("¿El target conserva memoria temporal?", "from statsmodels.tsa.stattools import acf; acf(monthly['cif_usd'],nlags=12)[[1,12]]",
  f"La autocorrelación es {num(r['acf_1'],4)} en t-1 y {num(r['acf_12'],4)} en t-12.",
  "Los rezagos recientes son esenciales y el rezago anual debe conservarse como referencia estacional."),

q("¿Cuáles son los estadísticos descriptivos de las variables monetarias y de peso?", "df[['PNK','VAFODO','FLETE','VACID','SEGUROS']].describe().T",
  "Por registro, la mediana de PNK es 757,67 kg y la de VACID es US$ 4.977,00; las medias son mucho mayores por la cola derecha.",
  "Para el nivel de registro se prefieren medianas, cuantiles y transformaciones robustas, no solo medias."),
q("¿Cómo son las distribuciones de las variables numéricas?", "df[['PNK','VAFODO','FLETE','VACID','SEGUROS']].quantile([.01,.25,.5,.75,.99])",
  "Las variables de peso, valor, flete y seguros presentan colas largas y diferencias marcadas entre mediana y extremos.",
  "El modelado mensual debe usar sumas y razones estables; si se modela por segmento, se aplicará log1p y winsorización solo dentro del train."),
q("¿Qué asimetría y curtosis presentan las variables numéricas?", "pd.DataFrame({'skew':df[['PNK','VAFODO','FLETE','VACID','SEGUROS']].skew(),'kurt':df[['PNK','VAFODO','FLETE','VACID','SEGUROS']].kurt()})",
  "En registros, VACID tiene asimetría 98,8494; FLETE 1.222,8707 y SEGUROS 2.304,9185.",
  "Los modelos lineales no deben recibir estos valores crudos sin transformación o agregación."),
q("¿Cuántos outliers detecta IQR en variables de registro?", "{c:((df[c]<df[c].quantile(.25)-1.5*(df[c].quantile(.75)-df[c].quantile(.25)))|(df[c]>df[c].quantile(.75)+1.5*(df[c].quantile(.75)-df[c].quantile(.25)))).sum() for c in ['PNK','VACID','FLETE']} ",
  "IQR marca 875.016 pesos netos, 571.693 valores CIF y 664.484 fletes.",
  "IQR no debe usarse como regla de borrado en comercio exterior; servirá para diagnóstico y transformaciones robustas."),
q("¿Qué outliers mensuales detecta z-score?", "s=monthly['cif_usd']; monthly.index[np.abs(stats.zscore(s))>3]",
  "Solo marzo y agosto de 2022 superan tres desviaciones estándar en la serie mensual.",
  "Se conservarán y se medirá el error del modelo en esos meses para evaluar robustez ante picos."),
q("¿Existen escalas numéricas muy dispares?", "df[['SEGUROS','FLETE','PNK','VACID','VACIP']].mean().sort_values()",
  f"La razón entre la mayor y menor media de las variables analizadas es aproximadamente {num(r['scale_ratio'],2)}.",
  "Ridge requiere StandardScaler ajustado solo con train; los modelos de árboles no requieren escalado."),
q("¿Hay ceros excesivos?", "df[['PNK','VAFODO','FLETE','VACID','TOTALIVAYO','SEGUROS','DEREL']].eq(0).mean().mul(100)",
  "DEREL contiene 46,232 % de ceros y TOTALIVAYO 6,072 %; en las demás variables principales la proporción es inferior a 0,2 %.",
  "Se crearán indicadores binarios de arancel cero y tributo cero cuando se modele composición; no se imputarán esos ceros."),
q("¿Qué transformaciones numéricas son razonables?", "np.log1p(df[['PNK','VAFODO','FLETE','VACID','SEGUROS']]).skew()",
  "log1p reduce las colas y acepta ceros. Box-Cox no es directamente aplicable a variables con cero y aportaría menos interpretabilidad.",
  "Se usará log1p para target y montos rezagados; porcentajes y razones se limitarán con cuantiles aprendidos en train."),
q("¿Qué variables mensuales están asociadas con CIF?", "monthly.select_dtypes('number').corr(method='spearman')['cif_usd'].sort_values(ascending=False)",
  "FOB, flete y seguros mensuales muestran correlaciones de 0,9971, 0,9066 y 0,7731 con CIF; peso neto alcanza 0,5181.",
  "Estas relaciones sirven para features rezagadas; las versiones del mismo mes se excluyen por fuga."),
q("¿La serie es estacionaria?", "from statsmodels.tsa.stattools import adfuller; adfuller(monthly['cif_usd'])[1], adfuller(np.log1p(monthly['cif_usd']).diff().dropna())[1]",
  f"ADF en nivel da p={r['adf_level_p']:.6f}; en primera diferencia logarítmica da p<{0.000001:.6f}.",
  "La serie en nivel no es estacionaria. ARIMA/SARIMA requeriría diferenciación; modelos supervisados usarán tendencia y rezagos."),

q("¿Qué países concentran el mayor valor CIF?", "df.groupby('PAISGEN')['VACID'].sum().nlargest(10)",
  "El código 215 lidera con US$ 66,18 mil millones; le siguen 493 y 249 con US$ 18,89 y US$ 18,75 mil millones.",
  "Se deben crear participaciones rezagadas de los países principales y agrupar el resto."),
q("¿Qué peso tienen las categorías raras de país?", "vc=df['PAISGEN'].value_counts(); vc[vc<100].sum()/len(df)*100",
  f"Los países con menos de 100 registros representan apenas {num(r['rare_country_share_pct'],3)} % de las filas.",
  "Las categorías raras pueden agruparse en 'Otros' sin perder volumen significativo."),
q("¿Existe alta cardinalidad categórica?", "df[['PAISGEN','capitulo','NABAN']].nunique()",
  f"País tiene {r['countries']} categorías y capítulo {r['chapters']}; subpartida es mucho más granular.",
  "Se evitará one-hot de subpartida. Para modelos globales se usarán capítulo, top-N y estadísticas históricas rezagadas."),
q("¿Qué capítulos arancelarios concentran el valor?", "df.groupby('capitulo')['VACID'].sum().nlargest(10)",
  "Los capítulos 85, 84, 87, 39 y 10 son los cinco principales; juntos explican 39,692 % del CIF.",
  "Se crearán participaciones mensuales rezagadas para estos capítulos y una categoría residual."),
q("¿Qué relación tienen las categorías principales con el target?", "shares=df.groupby('PAISGEN')['VACID'].sum().sort_values(ascending=False); shares.head(5).sum()/shares.sum()*100",
  "Los cinco países principales concentran 68,047 % del CIF, mientras los cinco capítulos principales concentran 39,692 %.",
  "La composición por país tiene mayor poder potencial, pero las participaciones deben calcularse con meses anteriores."),
q("¿Qué estrategia de encoding conviene?", "top=df['PAISGEN'].value_counts().head(15).index; pd.Series(np.where(df['PAISGEN'].isin(top),df['PAISGEN'].astype(str),'OTROS')).nunique()",
  "Un top-15 más 'Otros' reduce país a 16 categorías; para capítulo puede emplearse top-20 más 'Otros'.",
  "El one-hot se ajustará dentro del pipeline. Target encoding solo sería válido con codificación out-of-fold y respetando el tiempo."),

q("¿Cuál es la matriz de correlación mensual?", "monthly[['cif_usd','fob_usd','peso_neto_kg','flete_usd','seguros_usd','registros']].corr(method='spearman')",
  "CIF se mueve casi en paralelo con FOB y presenta asociación moderada con registros y peso.",
  "La matriz confirma qué variables deben rezagarse y ayuda a evitar duplicar señales contemporáneas."),
q("¿Qué rezagos se correlacionan más con el target?", "model_df.corr(method='spearman')['target'].sort_values(ascending=False).head(10)",
  "La media móvil CIF de 3 meses correlaciona 0,8501; CIF t-1, 0,8280; CIF t-2, 0,7838 y CIF t-3, 0,7735.",
  "El núcleo del modelo debe priorizar rezagos recientes y medias móviles calculadas con shift(1)."),
q("¿Existe multicolinealidad entre rezagos?", "pd.Series(vif).sort_values(ascending=False)",
  "Todos los VIF evaluados son altos; las medias móviles de 3 y 6 meses superan 2.000 por combinar los mismos rezagos.",
  "Ridge puede regularizar; en modelos lineales explicativos se seleccionará una sola ventana o se aplicará PCA."),
q("¿Qué interacciones pueden aportar señal?", "monthly.assign(flete_pct=monthly.flete_usd/monthly.fob_usd,precio=monthly.cif_usd/monthly.peso_neto_kg)[['flete_pct','precio']].describe()",
  "El costo logístico relativo y el precio implícito conectan valor, peso y flete de forma interpretable.",
  "Se crearán flete/FOB, CIF/kg y sus cambios, siempre rezagados para el horizonte t+1."),
q("¿Qué variables son redundantes?", "monthly[['cif_usd','fob_usd','flete_usd','seguros_usd']].corr(method='spearman')",
  "FOB contemporáneo es casi redundante con CIF (0,9971); las distintas medias móviles también son redundantes entre sí.",
  "Se limitará la duplicación de ventanas y se comparará el rendimiento por ablación."),
q("¿Qué variables muestran poco poder predictivo individual?", "model_df.corr(method='spearman')['target'].abs().sort_values().head()",
  "mes_cos tiene correlación -0,0056; peso t-12, 0,1137; y mes_sin, -0,1259 en términos absolutos.",
  "No se eliminarán solo por correlación: pueden aportar de forma no lineal, pero deberán justificar su permanencia en validación."),

q("¿Existe una tendencia de largo plazo?", "stats.linregress(np.arange(len(monthly)),monthly['cif_usd'])",
  f"La pendiente lineal es {usd(r['trend_slope_monthly'])} por mes, con R²={num(r['trend_r2'],4)} y p<0,001.",
  "Se incluirá índice temporal o tendencia, y se vigilará que no extrapole sin límites."),
q("¿Existe estacionalidad mensual?", "monthly.groupby(monthly.index.month)['cif_usd'].mean()/monthly['cif_usd'].mean()",
  "Agosto tiene el índice medio más alto, 1,0709; junio el más bajo, 0,9469. La amplitud media es moderada.",
  "Se usarán seno/coseno del mes y rezago 12; la estacionalidad se validará por ventanas móviles."),
q("¿Hay drift entre el periodo histórico y el reciente?", "stats.ks_2samp(monthly.loc[:'2021-12','cif_usd'],monthly.loc['2022-01':,'cif_usd'])",
  f"KS={num(r['drift_ks'],4)} con p<0,001; la mediana de 2022-2024 es {num(r['median_shift_pct'],3)} % mayor que la de 2012-2021.",
  "La validación debe dar mayor peso a ventanas recientes y el modelo debe reentrenarse al incorporar nuevos meses."),
q("¿Qué segmentos deben vigilarse por separado?", "df.groupby('PAISGEN')['VACID'].sum().nlargest(5), df.groupby('capitulo')['VACID'].sum().nlargest(5)",
  "Los países 215, 493 y 249 dominan el origen; los capítulos 85, 84 y 87 lideran productos.",
  "El producto debe mostrar pronóstico total y contribuciones de estos segmentos, con alerta cuando cambie su participación."),

q("¿Qué features finales se proponen?", "feature_set=['cif_lag_1','cif_lag_2','cif_lag_3','cif_lag_6','cif_lag_12','peso_lag_1','peso_lag_3','flete_lag_1','registros_lag_1','media_3','media_12','mes_sin','mes_cos','tendencia']; feature_set",
  "Se proponen rezagos CIF 1, 2, 3, 6 y 12; rezagos de peso, flete y registros; medias 3 y 12; calendario, tendencia y participaciones top-N rezagadas.",
  "El conjunto final se decidirá por validación walk-forward y ablación, no por correlación aislada."),
q("¿Cuál debe ser el pipeline de preprocesamiento?", "pipeline=['validar_mes_completo','normalizar_decimal','agregar_mes','crear_rezagos_con_shift','imputar_train','escalar_train','modelo']; pipeline",
  "El pipeline debe validar integridad, normalizar formatos, agregar por mes, crear rezagos, imputar y escalar dentro de cada fold.",
  "Toda transformación que aprenda parámetros debe ajustarse solo con entrenamiento para impedir fuga."),
q("¿Cómo debe hacerse el split y la validación?", "train=monthly.loc[:'2022-12']; test=monthly.loc['2023-01':'2024-12']; len(train),len(test)",
  "El backtest final usa 132 meses hasta diciembre de 2022 para train y 24 meses de 2023-2024 para prueba; dentro de train se aplicarán cortes expanding-window.",
  "No se usará train_test_split aleatorio. Cada predicción debe construirse únicamente con meses anteriores."),
q("¿Qué modelos y métricas resultan adecuados?", "pd.DataFrame(backtest).T.sort_values('WAPE_pct')",
  "En el holdout 2023-2024, Ridge logró MAE US$ 97,00 millones, RMSE US$ 118,48 millones, sMAPE 7,251 % y WAPE 7,223 %. El naive estacional tuvo WAPE 25,235 %.",
  "Ridge es el candidato inicial por precisión e interpretabilidad; se comparará con SARIMA y boosting. Las métricas principales serán WAPE y MAE, con RMSE como penalización de errores grandes."),
]

assert len(questions) == 52


def set_cell_margins(section):
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.75)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instruction)
    run._r.addnext(fld)


def set_font(run, bold=None, italic=None):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_labeled(doc, label, text, bold_label=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.08
    a = p.add_run(label)
    set_font(a, bold=bold_label)
    b = p.add_run(text)
    set_font(b)
    return p


doc = Document()
for section in doc.sections:
    set_cell_margins(section)

styles = doc.styles
for style in styles:
    if hasattr(style, "font"):
        style.font.name = "Arial"
        style.font.size = Pt(12)
        style.font.color.rgb = RGBColor(0, 0, 0)
for name in ["Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"]:
    st = styles[name]
    st.font.name = "Arial"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    st.font.size = Pt(12)
    st.font.color.rgb = RGBColor(0, 0, 0)
styles["Title"].font.bold = True
styles["Heading 1"].font.bold = True
styles["Heading 2"].font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run("Análisis exploratorio de datos de las importaciones registradas por la Aduana de Buenaventura"), bold=True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run("EDA de 52 preguntas orientado al pronóstico mensual del valor CIF"))
doc.add_paragraph()
add_labeled(doc, "Fuente principal: ", "DANE, Estadísticas de Importaciones IMPO, 2012-2024.")
add_labeled(doc, "Unidad temporal del modelo: ", "Mes.")
add_labeled(doc, "Objetivo: ", "Valor CIF mensual en dólares, con horizonte de un mes.")
doc.add_page_break()

doc.add_heading("Índice general", level=1)
toc = doc.add_paragraph()
add_field(toc, 'TOC \\o "1-2" \\h \\z \\u')
doc.add_page_break()

doc.add_heading("Ajustes al prompt original", level=1)
add_labeled(doc, "Adaptación aplicada: ", "Se conservaron las 52 preguntas, pero el análisis se ajustó a una serie temporal de regresión. Se sustituyó el enfoque de clases por distribución continua, se definió un horizonte t+1, se incorporaron estacionariedad, autocorrelación, drift y validación walk-forward, y se evaluó explícitamente la fuga temporal.")
add_labeled(doc, "Criterio de validez: ", "Los resultados provienen de la ejecución sobre 5.625.947 registros filtrados por ADUA=35. Los archivos oficiales se consolidaron respetando diferencias de codificación, delimitador y formato decimal.")

doc.add_heading("Preparación común del análisis", level=1)
setup_code = """import numpy as np
import pandas as pd
from scipy import stats
from statsmodels.tsa.stattools import adfuller, acf

ruta = 'data/landing/impo_buenaventura_2012_2024.csv.gz'
df = pd.read_csv(ruta, low_memory=False)
numericas = ['FECH','ADUA','PAISGEN','PAISPRO','PAISCOM','DEPTODES','VIATRANS','PBK','PNK','CANU','NABAN','VAFODO','FLETE','VACID','VACIP','VADUA','BASEIVA','TOTALIVAYO','SEGUROS','TIPOIM','PORARA','DEREL']
for c in numericas:
    s = df[c].astype('string').str.strip()
    local = s.str.replace('.', '', regex=False).str.replace(',', '.', regex=False)
    df[c] = pd.to_numeric(s.where(~s.str.contains(',', na=False), local), errors='coerce')
f = df['FECH'].round().astype('Int64')
df['fecha'] = pd.to_datetime({'year':2000+f//100,'month':f%100,'day':1}, errors='coerce')
df['capitulo'] = (df['NABAN']//100_000_000).astype('Int64')
monthly = df.groupby('fecha').agg(cif_usd=('VACID','sum'),fob_usd=('VAFODO','sum'),peso_neto_kg=('PNK','sum'),flete_usd=('FLETE','sum'),seguros_usd=('SEGUROS','sum'),registros=('FECH','size'),paises_origen=('PAISGEN','nunique'),capitulos=('capitulo','nunique')).sort_index().asfreq('MS')
monthly['precio_implicito_usd_kg'] = monthly['cif_usd']/monthly['peso_neto_kg'].replace(0,np.nan)"""
p = doc.add_paragraph()
for line in setup_code.splitlines():
    run = p.add_run(line + "\n")
    set_font(run)
add_labeled(doc, "Nota: ", "Los bloques de código de cada pregunta se ejecutan después de esta preparación común.")

groups = {
    1: "Estructura del dataset",
    7: "Calidad de datos",
    15: "Variable objetivo",
    23: "Variables numéricas",
    33: "Variables categóricas",
    39: "Relaciones y correlaciones",
    45: "Dimensión temporal y segmentación",
    49: "Conclusiones para el modelado",
}
fig_after = {16: "02_distribucion_target.png", 33: "05_paises.png", 39: "03_correlaciones.png", 45: "01_serie_cif.png", 46: "04_estacionalidad.png"}
fig_no = 0
for i, item in enumerate(questions, start=1):
    if i in groups:
        doc.add_page_break()
        doc.add_heading(groups[i], level=1)
    doc.add_heading(f"Pregunta {i}. {item['title']}", level=2)
    code_label = add_labeled(doc, "Código Python", "")
    code_label.paragraph_format.keep_with_next = True
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_together = True
    for line in item["code"].splitlines():
        set_font(p.add_run(line + "\n"))
    add_labeled(doc, "Respuesta o hallazgo: ", item["finding"])
    add_labeled(doc, "Implicación para el modelo: ", item["implication"])
    if i in fig_after:
        fig_no += 1
        doc.add_picture(str(FIGURES / fig_after[i]), width=Inches(6.2))
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(cp.add_run(f"Figura {fig_no}. Resultado asociado a la pregunta {i}."), italic=True)

doc.add_heading("Resumen ejecutivo", level=1)
summary = [
    "1. La base final reúne 5.625.947 registros de Buenaventura y 156 meses completos entre 2012 y 2024.",
    "2. Se corrigieron cambios reales de separador y formato decimal; ignorarlos habría creado ceros y rupturas falsas.",
    "3. El target CIF mensual es no estacionario en nivel, tiene autocorrelación t-1 de 0,8712 y drift reciente significativo.",
    "4. FOB y flete contemporáneos causan fuga; solo deben entrar rezagados.",
    "5. Los cinco países principales concentran 68,047 % del CIF y los cinco capítulos, 39,692 %.",
    "6. Se recomiendan rezagos, medias móviles desplazadas, calendario, tendencia y participaciones históricas por segmento.",
    "7. La validación debe ser temporal, con ventanas expansivas y holdout 2023-2024.",
    "8. Ridge alcanzó WAPE 7,223 %, frente a 25,235 % del naive estacional.",
    "9. El siguiente paso es repetir el backtest walk-forward y comparar Ridge, SARIMA y boosting.",
    "10. La inferencia debe bloquear meses incompletos y reentrenar cuando ingresen nuevos datos oficiales.",
]
for line in summary:
    p = doc.add_paragraph()
    set_font(p.add_run(line))

doc.add_heading("Fuentes", level=1)
sources = [
    "DANE. Estadísticas de Importaciones IMPO 2012-2024. https://microdatos.dane.gov.co/index.php/catalog/473/get-microdata",
    "DANE. Diccionario de datos de IMPO. https://microdatos.dane.gov.co/index.php/catalog/473/data_dictionary",
    "DANE. Código de aduana ADUA=35, Buenaventura. https://microdatos.dane.gov.co/catalog/473/datafile/F10/V385",
    "DIAN. Consultor de importaciones y exportaciones para seccionales. https://www.dian.gov.co/dian/cifras/Paginas/Consultor-de-Importaciones-y-Exportaciones-para-Seccionales.aspx",
]
for source in sources:
    p = doc.add_paragraph()
    set_font(p.add_run(source))

for section in doc.sections:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("Página "))
    add_field(footer, "PAGE")

docx_path = OUT / "EDA_52_Preguntas_Importaciones_Buenaventura.docx"
doc.save(docx_path)

nb = nbf.v4.new_notebook()
nb.metadata.kernelspec = {"display_name": "Python 3", "language": "python", "name": "python3"}
nb.cells.append(nbf.v4.new_markdown_cell("# EDA de 52 preguntas: importaciones por Buenaventura\n\nCuaderno ejecutable de respaldo del informe Word."))
nb.cells.append(nbf.v4.new_code_cell('%run "eda_compute.py"'))
for i, item in enumerate(questions, start=1):
    nb.cells.append(nbf.v4.new_markdown_cell(f"## Pregunta {i}. {item['title']}\n\n**Hallazgo:** {item['finding']}\n\n**Implicación:** {item['implication']}"))
    nb.cells.append(nbf.v4.new_code_cell(item["code"]))
notebook_path = ROOT / "reportes" / "soporte" / "EDA_52_Preguntas_codigo.ipynb"
notebook_path.parent.mkdir(parents=True, exist_ok=True)
if not notebook_path.exists():
    nbf.write(nb, notebook_path)

print(docx_path)
print(notebook_path)
